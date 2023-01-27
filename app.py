import streamlit as st
from streamlit_option_menu import option_menu
from docx import Document
from PIL import Image
import datetime
import dateTime
def downloadNow_faculty(data):
    file = Document()
    file.add_paragraph("{},".format(data["receiver"]))
    file.add_paragraph("Date:{}".format(data["starting_date"]))
    file.add_paragraph("Subject:{}".format(data["subject"]))
    file.add_paragraph("Sir/Ma'am")
    starting = data["starting_date"]
    end = data["end_date"]
    reason = data["reason"]
    no_of_day = dateTime.dateSubtract(end)
    person_name = data["person_name"]
    Name = data["Name"]
    file.add_paragraph("I am writing to ask you for a {} days leave from {} to {} due to {} . I am going to resume work from {}.".format(no_of_day,starting,end,reason,end))
    file.add_paragraph("I shall be reachable on my mobile number and email during the period. My person in charge, {} will be handling my tasks in my absence.".format(person_name))
    file.add_paragraph("""I will be thankful to you for considering my application.
    

Yours Sincerely,
    
{}""".format(Name))

    file.save("generatedLetter.docx")
    with open("generatedLetter.docx", "rb") as file:
        btn = st.download_button(
            label="Download Now",
            data=file,
            file_name="generatedLetter.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
          )
def home_faculty():
    name= st.text_input("Name")
    receiver = st.text_area("Receiver Details")
    subject = st.text_input("Enter your subject")
    date1 = str(st.date_input("Enter the starting date"))
    date2 = str(st.date_input("Enter the last date date"))
    reason = st.text_area("Enter The Reason ")
    person_name = st.text_input("Enter the name of the person who can represent you, during leave")
    if st.button("Submit"):
        datadict = {
        "Name" : name,
        "starting_date" : date1,
        "end_date" : date2,
        "receiver" : receiver,
        "subject" : subject,
        "reason" : reason,
        "person_name" : person_name
    }
        downloadNow_faculty(datadict)
def letter():
    st.title("Formal Application Generator")
    st.subheader("Dashboard")
    selected = option_menu(
        menu_title="Main Menu",
        options=["Home","Leave Application","About","Contact Us"],
        icons=["house","book","envelop","contact"],
        menu_icon="cast",
        orientation="horizontal"
        )
    if selected=="Home":
        st.markdown('''Hello
        EveryOne 
        
The formal letter generater is a tool that
can help quickly and easily construct a 
formal letter. 

Imagine this scenario: You need send a formal
to your employer or institution. This letter
needs to have the appropriate tone and style.

THANK YOU


''')
    if selected=="Leave Application":
        home_faculty()
    if selected == "Contact Us":
        st.markdown("""
        ## YOU CAN CONTACT ME -
        Instagram : mrtechnostart(Ram Badan Pandey)                  
        Email Id : rambpandey238@gmail.com
        """)
        st.text("IMS Engineering College, IMS Hostel, Ghaziabad Varanasi")
        st.markdown("""
        
        ## Team Members :-
        1. Ram Badan Pandey (2101430100141)
        2. Rohit Kumar Pandey (2101430100145)
        3. Saif Mohammad (2101430100150)
        4. Subodh Dubey (2101430100177)
        5. Suryansh Katiyar (2101430100180)
        """)
    if selected == "About":
        st.markdown('''We Are Here To Create Application ''')
def hideFooter():
    with Image.open("favicon.ico") as icon:
        st.set_page_config(page_title="ApploGen(beta)",page_icon=icon)
    hide_streamlit_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {
	            visibility: hidden;
            }
            footer:after {
                content:'Made With ❤️ By MrTechnoStart'; 
                visibility: visible;
                display: block;
                position: relative;
                #background-color: red;
                padding: 5px;
                top: 2px;
            }
                        </style>
                        """
    st.markdown(hide_streamlit_style, unsafe_allow_html=True)
hideFooter()
letter()